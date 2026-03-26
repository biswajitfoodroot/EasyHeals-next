import sys
import subprocess
import os

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set A4 margins
sections = doc.sections
for section in sections:
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# Title
title = doc.add_heading('EasyHeals Next: Current Functionality & Gap Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Analysis comparing current implementation vs HLDs, Architecture, and Roadmap Plans\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(text, level):
    h = doc.add_heading(text, level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(37, 99, 235) # TailWind Blue-600

def add_feature_list(features):
    for f in features:
        p = doc.add_paragraph(style='List Bullet')
        parts = f.split(":", 1)
        if len(parts) == 2:
            r = p.add_run(parts[0] + ":")
            r.bold = True
            p.add_run(parts[1])
        else:
            p.add_run(f)

# 1. Current Functionality Overview
add_heading('1. Current Implemented Functionality Overview', 1)
doc.add_paragraph("Based on the codebase analysis, the following modules are currently implemented and active:")

implemented_features = [
    "Authentication & Privacy: Fully functional OTP-based patient authentication (with Twilio integration), strict consent mechanisms per DPDP Act layout, Google OAuth for admins, and a soft-delete mechanism for 'right to erasure'.",
    "Data & Schema Models: Expansive SQLite schema managing Hospitals, Doctors, Appointments, Gamification records, Patient health profiles (EMR), Consultation rooms, and Pharmacy routing stubs.",
    "Patient Interaction APIs: Functioning routes for gamification tracking, leaderboard widgets, appointment booking interfaces (leads), EMR management, multi-lingual AI Intent search using Gemini.",
    "Admin & Moderation Dashboard: Extensive ingestion engine to discover and merge clinics dynamically via AI web scraping, content moderation tools for reviews and outlier data manipulation.",
    "Hospital & Doctor Profiles: Dynamic profiles generated with SEO enhancements (OpenGraph, Schema.org), integration of trust badge visibility, and detailed treatments rendering.",
    "Cron Processes: Outbox processes for external CRM synchronization and automated appointment reminders.",
]
add_feature_list(implemented_features)

# 2. Functionality Flow
add_heading('2. Core Functionality Flows', 1)

doc.add_paragraph("A brief walkthrough of the user flows as realized in the platform today:")

flows = [
    "Patient Discovery & Booking Flow: A patient lands on the SEO-optimized homepage (adaptive via geospatial IP matching), uses the Intent Search. They find a hospital or doctor, review trust badges, select a time, initiate a booking, and hit the DPDP Consent and OTP hurdle before the lead or CRM appointment is submitted.",
    "Administrative Data Ingestion: Administrators input URLs or search terms leading to background jobs (Ingestion Queue) which scrape provider details. The parsed data passes through AI merging logic and is staged in Candidate tables for human 'moderation' before going live into the Hospital/Doctor registry.",
    "Doctor/Hospital Routine: Verified hospital and doctor portal users sign in, manage availability (appointments queue), handle walk-ins or pending tele-consultations, and can access EMR summaries (from Gemini Vision) about the patient before visits.",
]
add_feature_list(flows)

# 3. Gap Identification
add_heading('3. Gap Identification vs HLD & Plan', 1)
doc.add_paragraph("By comparing the 'PHASE_SUMMARY.md' roadmap and 'HLD/Architecture' with the live code logic, the following differences and pending items (gaps) are surfaced:")

gaps = [
    "Phase 2 Complete vs Maturing (CRM Consolidation): While NextJS schemas integrate CRM fields, the 100% decoupling from the standalone Turso environment vs local CRM bridges may still have manual outbox sync gaps (Phase F). Messaging like WhatsApp Business is partly mapped but expects MSG91 fallback validation.",
    "Phase 3 Consultation Rooms (Paid vs Free Tiers): The schema has been built, and API routes ('/api/v1/consultations') are scaffolded. However, actual multi-participant RTC connections (Jitsi/Daily.co integrations with waiting rooms) heavily rely on external client-side hook integration which is still rolling out.",
    "Phase 4 ABHA/ABDM & Insurances (Missing): Integration with the National Health ID (ABHA) via the ABDM sandbox and Cashless Insurance validations are planned for Phase 4 but lack extensive explicit backend route definitions at the moment.",
    "Phase 5 Pharmacy Routing (Stubs Only): The schema shows stubs for 'pharmacies', 'patientFamilyLinks', and requests, but the active matching/routing pipelines for dispatching to Shadowfax/Porter delivery fleets are completely absent from 'src/app/api', aligning with a 'planned for later' status.",
    "Payment Integration Maturity: Razorpay integrations are staged in schema and routes but might require rigorous end-to-end webhook validation for real subscription handling (in hospital portals) out of sandbox environments."
]
add_feature_list(gaps)

# Conclusion
add_heading('4. Conclusion & Next Steps', 1)
doc.add_paragraph("The Phase 1 (Lead platform) and portions of Phase 2/3 (EMR + Consultation structures + Gamification APIs) are substantially complete in schema and backend representation. The immediate focus should likely be solidifying Phase 3's real-time video infrastructures or finishing the deeper Phase 2 CRM unifications, before initiating Phase 4 ABHA architectures. The application aligns strongly with the architectural directives defined in the project's markdown documentations.")

doc.save('Project_Functionality_And_Gaps.docx')
print("Document 'Project_Functionality_And_Gaps.docx' successfully generated.")
