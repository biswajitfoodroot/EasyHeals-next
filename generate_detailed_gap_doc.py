import sys
import subprocess
import os

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set A4 margins
sections = doc.sections
for section in sections:
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
title = doc.add_heading('EasyHeals Next: Comprehensive Functionality & Gap Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Exhaustive Code-Level Analysis of Functionality, Flows, and Comparisons to HLD\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(text, level):
    h = doc.add_heading(text, level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(37, 99, 235)

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# 1. Exhaustive List of Features
add_heading('1. Comprehensive Patient Functionality (Implemented Frontend & Backend)', 1)
doc.add_paragraph("Based on the /src/app routing and API architecture, the platform features a massive patient suite covering Clinical Continuity, Profiles, EMR, and Gamification:")

patient_apis = [
    ("ABHA Integration", "Route 'v1/patients/abha/link' handles the Indian National Health Authority ID linkage framework."),
    ("AI Health Coach", "Routes 'v1/patients/health-coach/conversations' integrate conversational AI models allowing patients to discuss symptoms safely."),
    ("AI Search & Intent", "Gemini-powered contextual search via 'v1/search/intent' allows patients to do natural language processing lookup for doctors, treatments, and locations."),
    ("EMR Modules", "APIs like 'v1/emr/lab-orders', 'prescriptions', 'visits', and 'vitals' allow patients to maintain and export ('v1/patients/health-export') a full longitudinal health record."),
    ("Document Sharing & Consent", "Patients upload documents ('v1/patients/documents'), use AI OCR ('prescription-scan'), and rigidly manage 'right to erasure' or revocation via 'v1/consent/revoke' and 'v1/patients/me'."),
    ("Patient Gamification", "A complete rewards loop tracking behavioral points, handled by 'v1/gamification/event' and 'v1/patients/rewards', tied into public leaderboards ('v1/leaderboard/[city]')."),
    ("Subscriptions & Payments", "Razorpay implementations mapped in 'v1/payments/membership/create-order' routing handle recurring or tier-based membership purchases."),
    ("Virtual Consultations", "Patients have interfaces at '/consultation/[sessionId]' connecting to 'v1/consultations/[id]/join' enabling RTC rooms with queue setups."),
    ("Patient UI Dashboards", "Includes specific screens for Records, Care Navigation, Upgrade, Privacy, Health Coach, and Timeline mapping ('/dashboard/health-timeline').")
]
for p in patient_apis:
    add_bullet(p[1], p[0])

add_heading('2. Comprehensive Provider (Doctor/Hospital) Functionality', 1)
doc.add_paragraph("Extensive standalone functionality exists within the /portal namespace for facility owners and practitioners:")

provider_apis = [
    ("Provider CRM Dashboards", "Separate portal pages exist for Hospitals ('/portal/hospital') and Doctors ('/portal/doctor') to view upcoming engagements and queue lists."),
    ("Appointment & Queue Live Tracking", "Routes 'v1/provider/queue/[id]/call' and 'v1/provider/schedule/generate' simulate live OPD room coordination and token progression."),
    ("Document Sharing Interlock", "Medical staff can access strictly shared patient document vaults via 'v1/portal/documents/shared' (protected by active consents)."),
    ("Pre-Visit Briefing", "AI endpoints like 'v1/previsit-briefs/[id]' auto-generate a comprehensive synopsis of the patient's EMR/timeline prior to the consultation directly for the doctor."),
    ("Staff & Entity Permissions", "Routes like 'v1/provider/staff' and 'portal/entity-permissions' manage role-based access for sub-users within an enterprise hospital setup."),
    ("Subscriptions & KYC", "Providers manage their own Tier-upgrades ('portal/subscription') and verifications ('portal/kyc-request') autonomously.")
]
for p in provider_apis:
    add_bullet(p[1], p[0])

add_heading('3. Comprehensive Administrative Data Ingestion & Moderation', 1)
doc.add_paragraph("The admin toolkit is massive and fully automated for generating the underlying directory:")

admin_apis = [
    ("AI Ingestion Pipeline", "Endpoints like 'admin/ingestion/jobs', 'discovery', and 'research-queue' manage crawling Google or hospital websites to parse doctor profiles, packages, and services."),
    ("Brochure Extractor", "Dedicated routes 'admin/research/brochure/scan' run AI Vision against uploaded rate-cards to automatically stage clinical packages and prices."),
    ("Moderation Queue", "All AI ingestion targets 'Candidates' tables; humans use 'v1/moderation/[id]/approve' or 'reject' to merge data safely into live 'hospitals' and 'doctors' tables."),
    ("System Configuration", "Full dynamic control system through 'admin/config/flags', 'admin/settings', allowing live swapping of AI features or notifications without redeployment."),
    ("Taxonomy Management", "APIs manage disease hierarchies via 'taxonomy/nodes', powering 'health-news', treatment routes ('/treatments/[slug]'), and the global search engine.")
]
for p in admin_apis:
    add_bullet(p[1], p[0])

add_heading('4. Detailed System Architecture & Functional Flow', 1)

flows = [
    ("Lead & Consult Flow (End-to-End)", "A patient searches via Intent Search -> Navigates to a canonical Hospital profile (e.g. '/hospital/[slug]') -> Chooses 'Book'. They log in via OTP (managed by 'register/send-otp'). They view the DPDP consent modal. Upon agreeing, their lead/appointment is persisted -> The Outbox Cron ('cron/outbox') synchronizes this lead back to the legacy CRM platform seamlessly."),
    ("Consultation Room Flow", "Once an appointment is active, the doctor triggers 'v1/consultations/[id]/start'. Patients connect via 'consultation/[sessionId]' and join the virtual waiting room. The provider calls them from the queue dashboard -> Post-visit, AI processes records into EMR schema logs."),
    ("Data Discovery Flow", "Admin queries 'Best cardiologists in Mumbai' into the research agent -> 'admin/ingestion/discovery' crawls SERP -> Data is mapped via Gemini into stubs -> Placed in 'ingestion_doctor_candidates'. An admin logs into '/admin', reviews the delta outlier score, and hits 'Approve' -> The doctor is now live in the global directory.")
]
for f in flows:
    add_bullet(f[1], f[0])

add_heading('5. Deep Gap Analysis vs HLD/Plan Documents', 1)
doc.add_paragraph("Cross-referencing the implemented API surface against 'PHASE_SUMMARY.md' & 'HLD':")

gaps = [
    ("Phase 1 & Phase 2 (Completed/Exceeding scope)", "Core CRM consolidation, lead routing, OTP infrastructure, and DPDP compliance limits exist natively. Furthermore, Gamification (expected Phase 2) is already highly structured via specific API modules and Dashboard routes, meaning development is ahead of schedule here."),
    ("Phase 3 (Clinical Continuity) - Maturing", "The EMR ecosystem is fully scaffolded ('lab-orders', 'visits', 'vitals'). Pre-visit Briefs are implemented ('v1/previsit-briefs'). The backbone for Consultation Rooms is built. GAP: Specific integrations with video SDK clients (Whereby/Jitsi) need further front-end binding to be production-ready."),
    ("Phase 4 (ABHA / Referral Networks) - Emerging", "The 'abha/link' route proves Phase 4 work has begun early. GAP: Insurance integration endpoints (Cashless check TPA bridges) are missing from the codebase."),
    ("Phase 5 (Pharmacy Routing) - Incomplete", "While the 'prescription-scan' endpoint allows users to OCR their prescriptions using AI, the actual quoting engines and delivery integrations (Shadowfax/Porter mapping mentioned in HLD) do not have corresponding functional routes. This is a true remaining gap aligned with the roadmap."),
    ("Notification Bridges", "The WhatsApp Business API / MSG91 pathways in NextJS rely largely on the cron outbox falling back to the legacy CRM platform to actually send messages, though stubs for Twilio natively exist.")
]
for g in gaps:
    add_bullet(g[1], g[0])

doc.save('Current_Functionality_Detailed.docx')
print("Document 'Current_Functionality_Detailed.docx' successfully generated.")
