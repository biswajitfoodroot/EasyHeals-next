import sys
import subprocess

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set A4 margins
sections = doc.sections
for section in sections:
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# Title
title = doc.add_heading('EasyHeals Next: Platform Functionalities', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Version 4.0 | Cross-Phase Architecture Reference (P1-P7) | Standalone Overview\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("This document outlines the key features and functionalities of the EasyHeals Next platform categorized by its four primary user personas: Patient, Hospital, Admin, and Doctor. The architecture supports AI-driven health navigation, deep EMR integration, gamification, and robust role-based access control.", style='Subtitle')

def add_heading(text, level):
    h = doc.add_heading(text, level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(37, 99, 235) # TailWind Blue-600

def add_feature_list(features):
    for f in features:
        p = doc.add_paragraph(style='List Bullet')
        parts = f.split(":", 1)
        if len(parts) == 2:
            r = p.add_run(parts[0] + ":")
            r.bold = True
            p.add_run(parts[1])
        else:
            p.add_run(f)

# Patient Features
add_heading('1. Patient Context & Features', 1)
patient_features = [
    "AI-Powered Intent Search: Translates natural language queries (even Hinglish) using Gemini into structured entity/specialty/location searches with FTS5 and Typesense.",
    "Conversational AI Health Coach: SSE streaming chat interface to engage patients, using contextual health memory (up to 200 previous events).",
    "Electronic Medical Records (EMR): Fully encrypted PHI-safe vault for visit records, lab orders, prescriptions, and vitals logging.",
    "Document Management & Sharing: Upload and manage medical documents (Vercel Blob) with granular, consent-gated sharing to providers.",
    "Gamification & Leaderboards: Streak badges, awards, phase-based milestones, and city-wide leaderboards for engagement and daily check-ins.",
    "Teleconsultation & Waiting Rooms: Jitsi-powered encrypted video consultation with waiting room and family/interpreter invites.",
    "Seamless Authentication & Privacy: OTP verification on WhatsApp/SMS, ABHA (Health ID) integration, strict DPDP consent management, and 'right to erasure' functionality.",
    "Real-time Appointment Booking: Four-step automated slot booking, real-time availability sync, and WhatsApp/SMS confirmation tracking.",
    "Payments & Memberships: Membership tiers integration using Razorpay for both subscription and one-time payments."
]
add_feature_list(patient_features)

# Hospital Features
add_heading('2. Hospital & Clinic Management Context', 1)
hospital_features = [
    "Role-Based Provider Portal: Dedicated dashboards for hospital admins, receptionists, and billing staff.",
    "Queue & Token Management: Live OPD queue and schedule progression visualization powered by Redis.",
    "Broadcasting & CRM Linking: MSG91-backed mass broadcast campaigns to segment out targeted messaging based on consent.",
    "Subscription & Self-registration: 4-step self-serve onboarding, tiered verifications, and Razorpay-powered facility subscriptions.",
    "Analytics Dashboard: Real-time insights covering appointment volumes, top doctors, slot utilization, and lead funnels.",
    "Staff Directory & Provisioning: Add, edit, and coordinate different providers and sub-users under one enterprise account.",
    "Automated SEO Profile Pages: Dynamically generated hospital and treatment profile pages with JSON-LD schemas and Breadcrumbs optimized for discoverability.",
    "AI Moderation of Data: Outdated profiles are automatically flagged via a staleness scan allowing 1-click confirmation by hospital administrators."
]
add_feature_list(hospital_features)

# Admin Features
add_heading('3. EasyHeals Administrator Context', 1)
admin_features = [
    "AI Scraping & Ingestion Pipeline: Intelligent web and brochure data scraping, automated staging and comparison tools for adding clinics/providers.",
    "Data Conflict Resolution & Approval Queue: A dedicated moderation dashboard linking confidence scores per ingested field, supporting approval, merge, or rejection.",
    "Rate Limiting & Threat Controls: Configuration panel adjusting DDoS protections, honeypots, rate ceilings, and bot detection settings continuously.",
    "Feature Flag Engine: Master control interface to toggle gamification, EMR, analytics, routing engines, or localized language testing without redeploying code.",
    "Data Privacy Auditing: Administrative capability to scrutinize token TTL, view anonymized activity trails, and govern consent logs.",
    "Content Moderation: Monitor leaderboard aliases, review provider profiles, and curate taxonomy tags, specialty synonyms, and master data structures.",
    "Global Outbox Health: Monitor cross-service communication to the legacy CRM pipeline via robust event-bus dashboards."
]
add_feature_list(admin_features)

# Doctor Features
add_heading('4. Doctor Context & Features', 1)
doctor_features = [
    "AI Pre-Visit Brief: Automatically generates concise clinical summaries for the doctor prior to the patient encounter via securely parsed document intelligence.",
    "Post-Visit Workflows: Quick form capture for updating patient EMRs with encrypted consultation notes, treatments, and prescriptions.",
    "Personalized Schedule View: Direct oversight of allocated daily slots, with the ability to pause intake or seamlessly inject walk-in/tele-consult shifts.",
    "Direct Lab Ordering: Prescribe lab tests tied directly to the hospital's internal integration, avoiding physical slips and increasing precision.",
    "Virtual Telehealth Desk: Control patient admission from virtual waiting rooms directly into active Jitsi sessions, capturing vital signs contemporaneously.",
    "Document Intelligence: Extract actionable data from uploaded lab charts (Gemini Vision integration) directly into structured EHR formats.",
    "Verified Profiles & Badging: Earn verifiable credentials and reviews boosting rank in regional searches."
]
add_feature_list(doctor_features)

doc.save('EasyHeals_Next_Features.docx')
print("Document 'EasyHeals_Next_Features.docx' successfully generated.")
